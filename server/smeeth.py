from docx import Document
from docx.shared import Inches
from openai import OpenAI
import os
import openai
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
import numpy as np

from docx import Document

from docx.shared import RGBColor
import re

client = OpenAI(api_key=openai.api_key)
doc = Document()


def convert_markdown_to_docx(markdown_file, output_docx):
    """Convert a Markdown file to a DOCX file."""
    # Read the Markdown file
    with open(markdown_file, "r") as file:
        lines = file.readlines()

    # Process each line from the Markdown file
    for line in lines:
        line = line.strip()  # Remove extra whitespace

        # Handle headers
        if line.startswith("# "):  # H1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):  # H3
            doc.add_heading(line[4:], level=3)

        # Handle bullet points
        elif line.startswith("- "):
            bullet_text = line[2:]  # Remove the bullet indicator

            # Check for the presence of bold text (indicated by **)
            if "**" in bullet_text:
                # Extract the bold text without bullet
                parts = bullet_text.split("**")

                paragraph = doc.add_paragraph()  # Create a new paragraph
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Odd parts are bold
                        run = paragraph.add_run(part.strip())  # Add bold text
                        run.bold = True
                    else:  # Even parts are normal text
                        paragraph.add_run(part.strip())  # Add normal text
                continue  # Skip adding this as a bullet point

            # If no bold formatting, just add it as a bullet point
            doc.add_paragraph(bullet_text, style="List Bullet")

        # Handle bold text (using ** for bold)
        elif "**" in line:
            paragraph = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd parts are bold
                    run = paragraph.add_run(part.strip())  # Add bold text
                    run.bold = True
                else:  # Even parts are normal text
                    paragraph.add_run(part.strip())  # Add normal text

        # Normal text
        elif line:
            doc.add_paragraph(line)

    # Save the Word document
    doc.save(output_docx)
    print(f"Markdown converted to Word file: {output_docx}")


# Example usage
convert_markdown_to_docx(
    "C:/Users/smeet.thadeshwar/Desktop/research-room-poc/reports/output_healthcare_uae_20241009_105152.md",
    "reports/output.docx",
)

########################################################################################################################

with open(
    "C:/Users/smeet.thadeshwar/Desktop/research-room-poc/reports/output_healthcare_uae_20241009_105152.md",
    "r",
) as file:
    content = file.read()


# Create a prompt for data extraction
prompt = (
    "Please extract numerical datapoints from the following Markdown content in Json format \n\n"
    f"{content}\n\n"
    "The data points are in the format 'Label: Value'."
)

# Send the prompt to the OpenAI API
response = client.chat.completions.create(
    model="gpt-4",  # Choose the correct model
    messages=[{"role": "user", "content": prompt}],
)

# Extracting the assistant's response
data_points_text = response.choices[
    0
].message.content  # Accessing the content correctly

# Step 2: Create a prompt for plotting
plot_prompt = f"""
Write a Python function that takes an input document and adds images of graphs to the report based on specific headings. The images are provided and should be inserted after the corresponding headings as follows:
 
1. After the **Executive Summary** section, insert the image `Market Size (USD Billion)_vs_Year_line_graph.png`.
 
2. After the **Sector Overview** section, insert the image `CAGR (%)_vs_Year_line_graph.png`.
 
3. After the **Competitive Landscape** section, insert the image `Revenue (USD Billion)_by_Company_bar_graph.png`.
 
Ensure that the images are centered and properly formatted for a professional report layout.
 
"""

response = client.chat.completions.create(
    model="gpt-4o",  # Choose the desired model
    messages=[{"role": "user", "content": plot_prompt}],
)


plot_code = response.choices[0].message.content

# For example, if the code has markdown syntax, you can clean it like this:
plot_code_clean = plot_code.replace("```python", "").replace("```", "")


heading_found = False


doc = Document("reports/output.docx")


# Variable to track if the heading is found
heading_found = False

# Dictionary mapping heading texts to image files
heading_to_image = {
    "Executive Summary": "Market Size (USD Billion)_vs_Year_line_graph.png",
    "Sector Overview": "CAGR (%)_vs_Year_line_graph.png",
    "Competitive Landscape": "Revenue (USD Billion)_by_Company_bar_graph.png",
}

# Variable to track if the heading is found
heading_found = False

# Iterate over paragraphs in the document
for index, paragraph in enumerate(doc.paragraphs):
    # Check each heading and insert corresponding image
    for heading, image in heading_to_image.items():
        # Check if the paragraph's style starts with 'Heading' and contains the heading text
        if paragraph.style.name.startswith("Heading") and heading in paragraph.text:
            # Insert a new paragraph before the heading to add the image
            img_paragraph = paragraph.insert_paragraph_before()
            img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the image

            # Add a run to insert the image
            run = img_paragraph.add_run()
            run.add_picture(image, width=Inches(4))  # Set width as needed

            heading_found = True
            print(f"Image {image} inserted before heading '{heading}'")
            break  # Break out of the inner loop once an image is inserted

if not heading_found:
    print("No headings found.")

# Save the changes to a new DOCX file
doc.save("updated_document2.docx")


##############################################################################################


def add_numbering_to_heading1(doc):
    """
    Adds numbering to only 'Heading 1' in the document, skipping 'Report Generated on:'.

    Args:
        doc: The Document object (from python-docx) where the headings need to be numbered.

    Returns:
        A list of paragraphs containing the numbered 'Heading 1' headings.
    """
    heading_count = 1  # Initialize heading numbering
    numbered_headings = []  # List to store the updated headings

    for paragraph in doc.paragraphs:
        # Check if the paragraph's style is exactly 'Heading 1'
        if paragraph.style.name == "Heading 1":
            # Skip the heading 'Report Generated on:'
            if "Report Generated on:" not in paragraph.text:
                # Add numbering to the heading text
                paragraph.text = f"{heading_count}. {paragraph.text}"
                numbered_headings.append(paragraph)  # Collect the numbered headings
                heading_count += 1  # Increment heading number

    return numbered_headings


# Example usage
doc = Document("updated_document2.docx")

# Call the function to add numbering to 'Heading 1' only, skipping specific heading
numbered_headings = add_numbering_to_heading1(doc)

# Save the document with numbered 'Heading 1' headings
doc.save("updated_document2.docx")
###################################################################################################################################################

# Load the document
doc = Document("updated_document2.docx")

# Define a regex pattern for matching URLs starting with https://
url_pattern = re.compile(r"https://[^\s]+")

# Iterate through each paragraph in the document
for para in doc.paragraphs:
    # Use the current paragraph text
    text = para.text
    # Find all matches in the paragraph text
    matches = url_pattern.findall(text)

    # If no matches, continue to the next paragraph
    if not matches:
        continue

    # Clear existing text
    para.clear()

    last_end = 0  # Track the end of the last match
    for url in matches:
        # Find the start index of the URL in the text
        start_index = text.find(url, last_end)

        # Add text before the URL
        if start_index > last_end:
            para.add_run(text[last_end:start_index])  # Add preceding text

        # Add the formatted URL run
        run = para.add_run(url)
        run.italic = True  # Italicize the URL
        run.font.color.rgb = RGBColor(0, 0, 255)  # Change text color to blue

        last_end = start_index + len(url)  # Update last_end to after the URL

    # Add any remaining text after the last URL
    if last_end < len(text):
        para.add_run(text[last_end:])  # Add remaining text

# Save the updated document
doc.save("updated_document_with_styled_urls.docx")
