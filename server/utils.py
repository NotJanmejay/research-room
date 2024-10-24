import os
import logging
from datetime import datetime
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_community.utilities import GoogleSerperAPIWrapper
import re
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import openai
from openai import OpenAI
import json
import matplotlib.pyplot as plt
import random
from docx.shared import Inches
from GraphPrompt import graph_prompt


# Load environment variables
def load_env():
    os.environ["REQUESTS_CA_BUNDLE"] = "./certificate.cer"
    load_dotenv()


def setup_logging(
    log_filename=f"logs/reports.log",
):
    if not os.path.exists("logs"):
        os.makedirs("logs")

    logging.basicConfig(
        filename=log_filename,
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


def initialize_model(model_id="gpt-4o-mini"):
    try:
        model = ChatOpenAI(model=model_id)
        logging.info("Initialized ChatOpenAI model.")
        return model
    except Exception as e:
        logging.error("Failed to initialize ChatOpenAI model: %s", e)
        print("Please add OPENAI_API_KEY to .env and try again")
        return None


def initialize_search(api_type="search"):
    try:
        if api_type == "news":
            search = GoogleSerperAPIWrapper(type="news")
        else:
            search = GoogleSerperAPIWrapper()
        logging.info(f"Initialized GoogleSerperAPIWrapper for {api_type}.")
        return search
    except Exception as e:
        logging.error("Failed to initialize GoogleSerperAPIWrapper: %s", e)
        print("Please add SERPER_API_KEY to .env and try again")
        return None


def fetch_company_details(domain, country, search, model):
    logging.info("Searching for top companies in the %s sector in %s.", domain, country)
    search_for_companies = search.results(
        f"Top Company in {domain} sector in {country}"
    )

    companies = model.invoke(
        f"Using the scrapped json value from google, give me the name of the top 5 companies working in the domain in comma separated value strictly only. Give me only the name of the companies and nothing else strictly, Scrapped Value: {search_for_companies}"
    )

    return companies.content.split(", ")


def fetch_company_financials(companies, search):
    companies_query = []

    for company in companies:
        logging.info("Fetching financial information for %s.", company)
        query_result = search.results(
            f"{company} market cap, revenue, debts, assets, financials"
        )
        companies_query.append({"company": company, "result": query_result})

    return companies_query


def fetch_market_insights(domain, country, search):
    prompts = {
        "market_size": f"Provide information about the current market size, historical trends, and future growth prospects of {domain} domain in {country}",
        "market_trends": f"Provide me detail key market trends in the field of {domain} in the country of {country}",
        "technological_advancement": f"Provide information about the latest technological advancements in the field of {domain} in {country}",
        "consumer_preferences": f"Provide information about the consumer preferences in the field of {domain} in the country of {country}",
        "regulatory_changes": f"Provide information about recent regulatory changes that took place in the field of {domain} in {country}",
        "barrier_to_entries": f"Provide information about high potential challenges or barriers to entries in the field of {domain} in {country}",
    }
    results = [
        search.results(prompts["market_size"]),
        search.results(prompts["market_trends"]),
        search.results(prompts["technological_advancement"]),
        search.results(prompts["consumer_preferences"]),
        search.results(prompts["regulatory_changes"]),
        search.results(prompts["barrier_to_entries"]),
    ]

    insights = {
        "market_size": results[0],
        "market_trends": results[1],
        "technological_advancement": results[2],
        "consumer_preferences": results[3],
        "regulatory_changes": results[4],
        "barrier_to_entry": results[5],
    }

    return insights


def fetch_latest_news(domain, country, search_news):
    latest_news = search_news.results(f"{domain} sector in {country}")
    return latest_news


def generate_report(
    domain, country, model, template, company_info, insights, latest_news
):
    relevant_information = (
        str(insights["market_size"])
        + str(insights["market_trends"])
        + str(insights["technological_advancement"])
        + str(insights["consumer_preferences"])
        + str(insights["regulatory_changes"])
        + str(insights["barrier_to_entry"])
    )

    latest_news_str = str(latest_news)
    prompt = template.format(
        domain=domain,
        country=country,
        current_year="2024",
        relevant_information=relevant_information,
        company_info=str(company_info),
        latest_news=latest_news_str,
    )

    logging.info("Generating report with the model.")
    report = model.invoke(prompt)
    return report.content


def save_report(content, domain: str, country: str):
    domain = "_".join(domain.strip().lower().split(" "))
    country = "_".join(country.strip().lower().split(" "))

    filename = f"reports_md/{domain}_{country}_{datetime.now().year}.md"

    if not os.path.exists("reports_md"):
        os.makedirs("reports_md")

    with open(filename, "w", encoding="utf-8") as file:
        file.write(content)

    return filename


def gpt_graph_generation(prompt):
    client = OpenAI()

    completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="gpt-4o-mini",
        response_format={"type": "json_object"},
    )

    logging.info("Generating json for Graphs")
    output = completion.choices[0].message.content
    logging.info("Graph JSON generated for the markdown")
    logging.info(output)
    return output


def generate_and_save_graphs(graphs):
    logging.info("Plotting and saving graphs to graph_images directory")
    bar_palettes = [
        [
            "teal",
            "steelblue",
            "skyblue",
            "dodgerblue",
            "deepskyblue",
            "lightskyblue",
            "powderblue",
            "cadetblue",
        ],
        ["#4D4D4D", "#5DA5DA", "#FAA43A", "#60BD68", "#F17CB0"],
        ["#A6CEE3", "#B2DF8A", "#FDBF6F", "#CAB2D6", "#FF7F00"],
    ]

    line_palettes = [
        ["teal", "seagreen", "#5DA5DA", "#FDBF6F", "#F17CB0"],
    ]

    pie_palettes = [
        plt.cm.Set3.colors,
        plt.cm.Accent.colors,
        ["#FFB347", "#FF6961", "#FDFD96", "#84B6F4", "#99C140"],
    ]

    graph_folder = "temp_graph_images"
    os.makedirs(graph_folder, exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    graph_paths = []

    for i, graph in enumerate(graphs):
        graph_type = graph["type"]
        heading = graph.get("heading", "Graph")
        x_values = graph["data"]["x_axis"]["values"]
        y_values = graph["data"]["y_axis"]["values"]
        x_label = graph["data"]["x_axis"]["label"]
        y_label = graph["data"]["y_axis"]["label"]

        if graph_type == "Bar Graph":
            colors = graph.get("colors", random.choice(bar_palettes))
        elif graph_type == "Line Graph":
            colors = graph.get("colors", random.choice(line_palettes))
        elif graph_type == "Pie Chart":
            colors = graph.get("colors", random.choice(pie_palettes))
        else:
            raise ValueError(f"Unsupported graph type: {graph_type}")

        plt.figure()

        if graph_type == "Bar Graph":
            plt.bar(x_values, y_values, color=colors)
        elif graph_type == "Line Graph":
            plt.plot(x_values, y_values, marker="o", color=colors[0])
        elif graph_type == "Pie Chart":
            plt.pie(y_values, labels=x_values, autopct="%1.1f%%", colors=colors)

        plt.title(heading)
        plt.xlabel(x_label)
        plt.ylabel(y_label)

        sanitized_heading = heading.replace(" ", "_")
        filename = f"{graph_folder}/{sanitized_heading}_{timestamp}_{i + 1}.png"
        plt.savefig(filename)
        graph_paths.append((heading, filename))
        plt.close()

    return graph_paths


def insert_graphs_for_heading(doc, graphs_to_insert, graph_paths):
    for path in graphs_to_insert:
        if os.path.exists(path):
            doc.add_paragraph()  # Add spacing before the graph
            doc.add_picture(path, width=Inches(4.5))
            doc.add_paragraph()  # Add spacing after the graph

            # Delete the graph after insertion
            try:
                os.remove(path)
                print(f"Deleted graph: {path}")
            except OSError as e:
                print(f"Error deleting {path}: {e}")


def process_text_with_hyperlinks(paragraph, text):
    source_pattern = r"\[Source:\s*([^\]]+)\]"
    parts = re.split(source_pattern, text)

    for i, part in enumerate(parts):
        if i % 2 == 1:
            domain = part.strip()
            url = f"https://{domain}"
            add_hyperlink(paragraph, f"Source: {domain}", url)
        else:
            part = part.strip()
            paragraph.add_run(part + ". ")

    url_matches = re.findall(r"(https?://[^\s]+)", text)
    for url in url_matches:
        if not re.search(source_pattern, text):
            add_hyperlink(paragraph, url, url)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")

    rPr.extend([color, underline])
    run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def convert_markdown_to_docx(markdown_content, file_name, graphs):
    if not os.path.isdir("reports_docx"):
        os.mkdir("reports_docx")
    """Convert markdown content to a DOCX file with proper formatting."""
    report_folder = "report"
    os.makedirs(report_folder, exist_ok=True)

    base_file_name = file_name.replace("reports/", "").replace(".md", "")
    base_output_docx = os.path.join(report_folder, base_file_name)

    output_docx = f"{base_output_docx}.docx"
    if os.path.exists(output_docx):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_docx = f"{base_output_docx}_{timestamp}.docx"

    doc = Document()

    lines = markdown_content.splitlines()
    graph_json = json.loads(gpt_graph_generation(graph_prompt + markdown_content))
    graph_paths = generate_and_save_graphs(graph_json["graphs"])

    current_heading = None
    added_graphs = set()  # Track added graphs for sections
    graphs_to_insert = []  # Store graphs to insert later

    is_references_section = False

    for line in lines:
        line = line.strip()

        if line.startswith("---"):
            continue

        if line in ["## 11. References Used", "### 11. References Used"]:
            is_references_section = True
            doc.add_heading(line[3:].strip(), level=2)
            doc.add_paragraph()
            continue

        if is_references_section:
            if not line:
                continue
            url_match = re.search(r"\[(.*?)\]\((https?://[^\s]+)\)", line)
            if url_match:
                paragraph = doc.add_paragraph(style="List Bullet")
                link_text, url = url_match.groups()
                add_hyperlink(paragraph, link_text, url)
            else:
                paragraph.add_run(line)

        if line.startswith("# "):
            if (
                current_heading
            ):  # Insert graphs for the previous section before moving to a new heading
                insert_graphs_for_heading(doc, graphs_to_insert, graph_paths)
            current_heading = line[2:].replace("*", "").strip()
            doc.add_heading(line[2:].replace("*", "").strip(), level=1)
            graphs_to_insert = []  # Reset the graph list for the new section
        elif line.startswith("## "):
            if (
                current_heading
            ):  # Insert graphs for the previous section before moving to a new heading
                insert_graphs_for_heading(doc, graphs_to_insert, graph_paths)
            current_heading = line[3:].replace("*", "").strip()
            doc.add_heading(line[3:].replace("*", "").strip(), level=2)
            graphs_to_insert = []  # Reset the graph list for the new section
        elif line.startswith("### "):
            if (
                current_heading
            ):  # Insert graphs for the previous section before moving to a new heading
                insert_graphs_for_heading(doc, graphs_to_insert, graph_paths)
            current_heading = line[4:].replace("*", "").strip()
            doc.add_heading(line[4:].replace("*", "").strip(), level=3)
            graphs_to_insert = []
        elif line.startswith("- "):
            bullet_text = line[2:]

            # Check for a source hyperlink at the end of the bullet point
            source_match = re.search(r"\[Source:\s*([^\]]+)\]", bullet_text)
            if source_match:
                domain = source_match.group(1).strip()
                bullet_text = bullet_text[
                    : source_match.start()
                ].strip()  # Remove the source part from the bullet text
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.add_run(bullet_text)  # Add the bullet text

                # Add the source as a hyperlink
                add_hyperlink(paragraph, f"Source: {domain}", f"https://{domain}")
            else:
                if "*" in bullet_text:
                    paragraph = doc.add_paragraph(style="List Bullet")
                    parts = bullet_text.split("**")
                    for i, part in enumerate(parts):
                        run = paragraph.add_run(part.strip())
                        run.bold = i % 2 == 1
                else:
                    hyperlink_match = re.match(
                        r"\[(.*?)\]\((https?://[^\s]+)\)", bullet_text
                    )
                    if hyperlink_match:
                        link_text, url = hyperlink_match.groups()
                        paragraph.add_run(" ")
                        add_hyperlink(paragraph, link_text, url)
                    else:
                        doc.add_paragraph(bullet_text, style="List Bullet")
        elif "**" in line:
            paragraph = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = paragraph.add_run(part.strip())
                run.bold = i % 2 == 1
        elif line:
            paragraph = doc.add_paragraph()
            process_text_with_hyperlinks(paragraph, line)

        # Store graph for the current section if it matches any graph heading
        for heading, path in graph_paths:
            if (
                heading.lower()
                == (re.sub(r"[^A-Za-z\s]", "", current_heading).strip()).lower()
                and heading not in added_graphs
            ):
                graphs_to_insert.append(path)  # Store the path to insert later
                added_graphs.add(heading)  # Mark this graph as added
                break  # Exit after adding the matching graph path
    # Insert graphs for the last section after processing all lines
    if current_heading:
        insert_graphs_for_heading(doc, graphs_to_insert, graph_paths)

    doc.save(output_docx)
    return output_docx


def convert_word_to_pdf(word_file_path):
    logging.info("Converting word file to pdf format")
    pdf_folder = "reports_pdf"
    os.makedirs(pdf_folder, exist_ok=True)

    pdf_path = os.path.join(
        pdf_folder, os.path.basename(word_file_path).replace(".docx", ".pdf")
    )

    convert(word_file_path, pdf_path)
    logging.info("Report generated for this session")

    return pdf_path
